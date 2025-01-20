import os
import requests
import streamlit as st
from langchain.schema import Document
from langchain_community.document_loaders import TextLoader
from embeddings import Embedder
from langchain_community.chat_models import ChatOpenAI
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import DocArrayInMemorySearch
from langchain.chains import ConversationalRetrievalChain
import os
os.environ["GIT_PYTHON_GIT_EXECUTABLE"] = r"C:\Program Files\Git\bin\git.exe"
import git
import io
from openai import OpenAI
from PIL import Image
import base64
from queue import Queue  
import traceback
import json
import datetime
from debug_logger import DebugLogger
from pathlib import Path
import base64


def count_tokens(text):
    """Count tokens accurately"""
    if not text:
        return 0
    
    print("\n=== Token Calculation Details ===")
    print(f"Text length: {len(text)} characters")
    
    # Separate by spaces and common punctuation
    words = text.split()
    print(f"Words count: {len(words)}")
    
    # Adjustment factors
    special_chars = sum(not c.isalnum() for c in text)
    print(f"Special characters: {special_chars}")
    
    numbers = sum(c.isdigit() for c in text)
    print(f"Numbers: {numbers}")
    
    # Base approach
    estimated_tokens = len(words) + (special_chars * 0.3) + (numbers * 0.5)
    print(f"Estimated tokens: {estimated_tokens}")
    
    return int(estimated_tokens)


class TokenCounter:
    def __init__(self):
        self.token_cache = {}
        
    def get_tokens(self, content_id, content):
        """Get tokens with caching"""
        print(f"\n=== TokenCounter.get_tokens ===")
        print(f"Content ID: {content_id}")
        
        if content_id in self.token_cache:
            cached_tokens = self.token_cache[content_id]
            print(f"Cache hit: {cached_tokens} tokens")
            return cached_tokens
            
        print("Cache miss - calculating tokens...")
        tokens = count_tokens(content)
        print(f"Calculated tokens: {tokens}")
        
        self.token_cache[content_id] = tokens
        print("Tokens cached")
        
        return tokens

class TokenUpdateObserver:
    def __init__(self):
        self.total_tokens = 0
        self.source_tokens = {
            'files': 0,
            'jira': 0,
            'confluence': 0,
            'github': 0
        }
        
    def update(self, source_type, tokens):
        """Update token counts when content changes"""
        self.source_tokens[source_type] += tokens
        self.total_tokens = sum(self.source_tokens.values())
        
        # Update UI
        st.session_state.token_stats = {
            'total': self.total_tokens,
            'breakdown': self.source_tokens.copy()
        }


class Utils:
    def __init__(self, st):
        self.docs = []
        if 'utils_docs' not in st.session_state:
            st.session_state.utils_docs = []
        self.jira_docs = []
        self.confluence_docs = []
        self.model = None
        self.hf = None
        self.db = None
        self.retriever = None
        self.MyQueue = Queue(maxsize=2)
        self.file_contents = {}
        self.st = st
        self.icons_path = Path(__file__).parent / 'icons'
        if not self.icons_path.exists():
            print(f"Warning: Icons directory not found at {self.icons_path}")
            # Try to create the icon directory if it doesn't exist
            try:
                self.icons_path.mkdir(parents=True, exist_ok=True)
                print(f"Created icons directory at {self.icons_path}")
            except Exception as e:
                print(f"Error creating icons directory: {str(e)}")
        # Adding the new instances
        self.token_counter = TokenCounter()
        self.token_observer = TokenUpdateObserver()
        
        # Restore previous counts if they exist
        if 'saved_token_counts' in st.session_state:
            self.token_counts = st.session_state.saved_token_counts
        else:
            self.token_counts = {
                'files': {},
                'jira': {},
                'confluence': {},
                'github': {}
            }
        
        # Initialize processed_files in session_state if not exists
        if 'processed_files' not in st.session_state:
            st.session_state.processed_files = set()
        
        # Initialize file_token_count in session_state if not exists
        if 'file_token_count' not in st.session_state:
            st.session_state.file_token_count = {}
        
        # Initialize token_stats to session_state if not exists
        if 'token_stats' not in st.session_state:
            st.session_state.token_stats = {
                'total': 0,
                'breakdown': {
                    'files': 0,
                    'jira': 0,
                    'confluence': 0,
                    'github': 0
                }
            }
        
        # Initialize attachments in session_state if not exists
        if 'attachments' not in st.session_state:
            st.session_state.attachments = {}
        
        # Initialize file_contents in session_state if not exists
        if 'file_contents' not in st.session_state:
            st.session_state.file_contents = {}
            
        # Initialize debug mode if it does not exist
        if 'debug_mode' not in st.session_state:
            st.session_state.debug_mode = False
            
        # Initialize log_to_file if not exists
        if 'log_to_file' not in st.session_state:
            st.session_state.log_to_file = True
        
        # Check requirements
        self.check_requirements()
        
        # Initialize logger
        self.debug = DebugLogger(
            enabled=st.session_state.get('debug_mode', False),
            log_to_file=st.session_state.get('log_to_file', True)
        )
        
        # Initialization confirmation print
        print("\n=== Utils Initialization Complete ===")
        print(f"- Documents loaded: {len(self.docs)}")
        print(f"- Token counter initialized: {bool(self.token_counter)}")
        print(f"- Token observer initialized: {bool(self.token_observer)}")
        print(f"- Session state keys: {list(st.session_state.keys())}")
        print(f"- Debug mode: {st.session_state.debug_mode}")
        print(f"- Log to file: {st.session_state.log_to_file}")
        
        # Log initialization
        if hasattr(self, 'debug'):
            self.debug.log("Utils initialized successfully", "info")

    def remove_file(self, file_path):
        """Delete file and its tokens consistently"""
        normalized_path = os.path.normpath(file_path)
        
        # Clear token counts
        if normalized_path in self.token_counts['files']:
            del self.token_counts['files'][normalized_path]
            
        # Clear session state
        if normalized_path in self.st.session_state.file_token_count:
            del self.st.session_state.file_token_count[normalized_path]
            
        #Clean docs
        self.docs = [doc for doc in self.docs if os.path.normpath(doc.metadata.get('source', '')) != normalized_path]
        
        # Clean processed files
        if normalized_path in self.st.session_state.processed_files:
            self.st.session_state.processed_files.remove(normalized_path)
            
        # Save state
        self.save_token_counts()
        
        print(f"Removed file: {normalized_path}")

    def prepare_image(self, image_path):

        """Prepare image to display without checkerboard pattern"""

        img = Image.open(image_path)
        if img.mode in ('RGBA', 'LA'):
            background = Image.new('RGBA', img.size, (0, 0, 0, 0))
            background.paste(img, mask=img.split()[-1])
            return background
        return img
    
    
    def count_tokens(self, text):
        """Count tokens accurately"""
        if not text:
            return 0
        
        # Print text details 
        print("\n=== Token Calculation Details ===")
        print(f"Text length: {len(text)} characters")
        
        # Split by spaces and common punctuation
        words = text.split()
        print(f"Words count: {len(words)}")
        
        # Adjustment factors
        special_chars = sum(not c.isalnum() for c in text)
        print(f"Special characters: {special_chars}")
        
        numbers = sum(c.isdigit() for c in text)
        print(f"Numbers: {numbers}")
        
        # Base approximation
        estimated_tokens = len(words) + (special_chars * 0.3) + (numbers * 0.5)
        print(f"Estimated tokens: {estimated_tokens}")
        
        return int(estimated_tokens)

    def display_file_info(self, file_path):
        normalized_path = os.path.normpath(file_path)
        tokens = self.token_counts['files'].get(normalized_path, 0)
        filename = os.path.basename(normalized_path)
        
        return f"{filename} ({tokens:,} tokens)"

    def save_token_counts(self):
        """Save counts in session_state """
        self.st.session_state.saved_token_counts = self.token_counts.copy()


    def process_file_tokens(self, file_path):
        try:
            normalized_path = os.path.normpath(file_path)
            print("\n=== Processing File Tokens ===")
            print(f"File: {file_path}")
                
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                print(f"Content length: {len(content)} characters")
                
            # Calculate tokens using the token_counter
            tokens = self.token_counter.get_tokens(normalized_path, content)
            print(f"Calculated tokens: {tokens}")
            
            # Update counts
            previous_tokens = self.token_counts['files'].get(normalized_path, 0)
            print(f"Previous token count: {previous_tokens}")
            
            self.token_counts['files'][normalized_path] = tokens
            print(f"Updated token count: {tokens}")
            
            # Update session_state
            if 'file_token_count' not in self.st.session_state:
                self.st.session_state.file_token_count = {}
            
            previous_session_tokens = self.st.session_state.file_token_count.get(normalized_path, 0)
            print(f"Previous session token count: {previous_session_tokens}")
            
            self.st.session_state.file_token_count[normalized_path] = tokens
            print(f"Updated session token count: {tokens}")
            
            #   Update observer
            self.token_observer.update('files', tokens)
            print(f"Observer updated with {tokens} tokens")
            
            return tokens
            
        except Exception as e:
            print(f"Error counting tokens for {file_path}: {str(e)}")
            traceback.print_exc()
            return 0
        
    def update_token_count(self, source_type, identifier, content):
        """Update token count with improved consistency"""
        try:
            #Normalize identifier
            if source_type == 'files':
                identifier = os.path.normpath(identifier)
                
            # Calculate tokens
            tokens = self.token_counter.get_tokens(identifier, content)
            
            # Actualizar conteos
            self.token_counts[source_type][identifier] = tokens
            
            # Update session state
            if 'file_token_count' not in self.st.session_state:
                self.st.session_state.file_token_count = {}
                
            if source_type == 'files':
                self.st.session_state.file_token_count[identifier] = tokens
                
            # Update total
            self.token_observer.update(source_type, tokens)
            
            print(f"Updated tokens for {identifier}: {tokens}")
            return tokens
            
        except Exception as e:
            print(f"Error updating token count: {str(e)}")
            return 0

    def get_total_tokens(self):
        """Get total tokens across all resources"""
        return sum(sum(counts.values()) for counts in self.token_counts.values())


    def process_screenshot(self, image_data):
        """Process screenshot data and save as attachment"""
        try:
            # Check if we are already processing an image
            if 'processing_image' in self.st.session_state:
                print("Already processing an image, skipping...")
                return None
                
            # Mark that we are processing an image 
            self.st.session_state.processing_image = True
            
            print("\n=== Starting Screenshot Processing ===")
            print(f"Image data type: {type(image_data)}")
            print("Memory state check:")
            print(f"- utils_docs in session: {'utils_docs' in self.st.session_state}")
            print(f"- selected_resources in session: {'selected_resources' in self.st.session_state}")
            print(f"- file_contents in session: {'file_contents' in self.st.session_state}")
            
            # Generate unique name for screenshot
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"screenshot_{timestamp}.png"
            normalized_path = os.path.normpath(filename)
            print(f"Generated filename: {filename}")
            print(f"Normalized path: {normalized_path}")
            
            # Convert image to bytes if necessary
            print("\n=== Image Conversion ===")
            try:
                if isinstance(image_data, Image.Image):
                    print("Converting PIL Image to bytes...")
                    img_byte_arr = io.BytesIO()
                    image_data.save(img_byte_arr, format='PNG')
                    img_bytes = img_byte_arr.getvalue()
                    width, height = image_data.size
                    print(f"Converted successfully - Size: {len(img_bytes)} bytes")
                    print(f"Dimensions: {width}x{height}")
                else:
                    img_bytes = image_data
                    print(f"Using raw image data - Size: {len(img_bytes)} bytes")
                    try:
                        temp_img = Image.open(io.BytesIO(img_bytes))
                        width, height = temp_img.size
                        print(f"Dimensions: {width}x{height}")
                    except Exception as e:
                        print(f"Could not determine dimensions: {str(e)}")
                        width = height = "unknown"
            except Exception as e:
                print(f"Error in image conversion: {str(e)}")
                if 'processing_image' in self.st.session_state:
                    del self.st.session_state.processing_image
                raise

            print("\n=== Creating Metadata ===")
            metadata = {
                'source': normalized_path,
                'normalized_path': normalized_path,
                'original_path': normalized_path,
                'original_filename': filename,
                'title': filename,
                'filename': filename,
                'type': 'files',
                'file_type': 'image',
                'timestamp': timestamp,
                'size': len(img_bytes),
                'width': width,
                'height': height,
                'processed_timestamp': datetime.datetime.now().isoformat(),
            }
            print("Metadata created successfully")
            print(f"Metadata keys: {list(metadata.keys())}")

            print("\n=== Processing Vision API ===")
            try:
                vision_result = self.process_image(image_data)
                print(f"Vision API result length: {len(vision_result) if vision_result else 0}")
            except Exception as e:
                print(f"Vision API error: {str(e)}")
                vision_result = f"Error analyzing image: {str(e)}"

            print("\n=== Creating Document Content ===")
            content = f"""
            Screenshot Information:
            Filename: {metadata['filename']}
            Timestamp: {metadata['timestamp']}
            Dimensions: {width}x{height}
            Size: {len(img_bytes)} bytes
            
            Image Analysis:
            {vision_result}
            """
            print(f"Content length: {len(content)}")

            # Calculate tokens
            tokens = self.token_counter.get_tokens(filename, content)
            print(f"\nToken Analysis:")
            print(f"- Content length: {len(content)} characters")
            print(f"- Calculated tokens: {tokens}")
            
            metadata.update({
                'content_length': len(content),
                'token_count': tokens
            })

            # Create document
            doc = Document(
                page_content=content,
                metadata=metadata
            )
            print("Document created successfully")

            print("\n=== Updating Session State ===")
            # Add to utils_docs (Local Files)
            if 'utils_docs' not in self.st.session_state:
                print("Initializing utils_docs in session state")
                self.st.session_state.utils_docs = []
            self.st.session_state.utils_docs.append(doc)
            print(f"utils_docs count: {len(self.st.session_state.utils_docs)}")

            # Add to local docs
            self.docs.append(doc)
            print(f"Local docs count: {len(self.docs)}")

            # Update token counters
            self.token_counts['files'][normalized_path] = tokens
            if 'file_token_count' not in self.st.session_state:
                self.st.session_state.file_token_count = {}
            self.st.session_state.file_token_count[normalized_path] = tokens
            print("Token counters updated")

            # Update selected resources for Local Files
            if 'selected_resources' not in self.st.session_state:
                print("Initializing selected_resources")
                self.st.session_state.selected_resources = {'files': {}}
            if 'files' not in self.st.session_state.selected_resources:
                print("Initializing files in selected_resources")
                self.st.session_state.selected_resources['files'] = {}
            self.st.session_state.selected_resources['files'][normalized_path] = {
                'selected': True,
                'type': 'image',
                'timestamp': timestamp
            }
            print("Selected resources updated")
            print(f"Current selected resources: {self.st.session_state.selected_resources['files'].keys()}")

            # Update file_contents to display the image
            if 'file_contents' not in self.st.session_state:
                print("Initializing file_contents")
                self.st.session_state.file_contents = {}
            self.st.session_state.file_contents[normalized_path] = {
                'content': content,
                'metadata': metadata,
                'data': img_bytes,
                'type': 'image'
            }
            print("File contents updated")
            print(f"Current file_contents keys: {self.st.session_state.file_contents.keys()}")

            #Add to processed_files
            if 'processed_files' not in self.st.session_state:
                self.st.session_state.processed_files = set()
            self.st.session_state.processed_files.add(normalized_path)
            print(f"Current processed files: {self.st.session_state.processed_files}")

            # Update embedder
            if hasattr(self.st.session_state, 'embedder'):
                print("Adding to embedder...")
                self.st.session_state.embedder.add_document(doc)
                print("Added to embedder successfully")

            print("\n=== Final Verification ===")
            print(f"- Document in utils_docs: {any(d.metadata['source'] == normalized_path for d in self.st.session_state.utils_docs)}")
            print(f"- Path in selected_resources: {normalized_path in self.st.session_state.selected_resources['files']}")
            print(f"- Path in file_contents: {normalized_path in self.st.session_state.file_contents}")
            print(f"- Path in processed_files: {normalized_path in self.st.session_state.processed_files}")

            print("\n=== Screenshot Processing Complete ===")

            # Clear processing flag before returning
            if 'processing_image' in self.st.session_state:
                del self.st.session_state.processing_image

            return {
                'success': True,
                'status': 'processed',
                'message': f'Screenshot processed successfully: {filename}',
                'metadata': metadata,
                'tokens': tokens,
                'filename': filename
            }
            
        except Exception as e:
            # Make sure to clear the processing flag in case of error
            if 'processing_image' in self.st.session_state:
                del self.st.session_state.processing_image
                
            print("\n=== Error in process_screenshot ===")
            print(f"Error type: {type(e).__name__}")
            print(f"Error message: {str(e)}")
            print("Full traceback:")
            traceback.print_exc()
            return {
                'success': False,
                'status': 'error',
                'message': f'Error processing screenshot: {str(e)}',
                'metadata': None
            } 


    def check_requirements(self):
        """Check if required packages are installed"""
        required_packages = {
            'python-docx': 'docx',
            'pandas': 'pandas',
            'PyPDF2': 'PyPDF2'
        }
        
        missing_packages = []
        
        for package, import_name in required_packages.items():
            try:
                __import__(import_name)
            except ImportError:
                missing_packages.append(package)
        
        if missing_packages:
            self.st.warning(f"Required packages missing: {', '.join(missing_packages)}. "
                          f"Please install them using:\n\n"
                          f"pip install {' '.join(missing_packages)}")
    

    def get_file_metadata(self, file_path):
        """
        Gets the contents and token count of a file from metadata
        """
        try:
            print(f"\n=== Getting file metadata for: {file_path} ===")
            normalized_path = os.path.normpath(file_path)
            
            # Search in utils_docs
            if hasattr(self.st.session_state, 'utils_docs'):
                for doc in self.st.session_state.utils_docs:
                    if os.path.normpath(doc.metadata.get('source', '')) == normalized_path:
                        content = doc.metadata.get('content')
                        token_count = doc.metadata.get('token_count', 0)
                        return {
                            'content': content,
                            'token_count': token_count,
                            'source': normalized_path
                        }
            
            # If not in utils_docs, check in token_counts
            token_count = self.get_file_tokens(normalized_path)
            
            # If exists in file_contents, get the contents
            content = None
            if 'file_contents' in self.st.session_state:
                content = self.st.session_state.file_contents.get(normalized_path)
            
            return {
                'content': content,
                'token_count': token_count,
                'source': normalized_path
            }
            
        except Exception as e:
            print(f"Error getting file metadata: {str(e)}")
            return None

    def get_file_tokens(self, file_path):
        """Get token count for a file"""
        try:
            normalized_path = os.path.normpath(file_path)
            
            # First look in token_counts
            if normalized_path in self.token_counts['files']:
                return self.token_counts['files'][normalized_path]
            
            # If not in token_counts, look in session_state
            if 'file_token_count' in self.st.session_state:
                return self.st.session_state.file_token_count.get(normalized_path, 0)
            
            return 0
            
        except Exception as e:
            print(f"Error getting file tokens: {str(e)}")
            return 0

    def handle_attachments(self, uploaded_file, allowed_types=None):
        """
        Process uploaded files with improved token and metadata handling.
        """
        try:
            print("\n=== Starting Attachment Processing ===")
            print(f"File Details:")
            print(f"- Name: {uploaded_file.name}")
            print(f"- Type: {uploaded_file.type}")
            print(f"- Size: {uploaded_file.size} bytes")
            
            # Generate unique name with timestamp
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            base_name = os.path.splitext(uploaded_file.name)[0]
            file_extension = os.path.splitext(uploaded_file.name)[1].lower()
            new_filename = f"{base_name}_{timestamp}{file_extension}"
            
            print(f"Generated filename: {new_filename}")

            try:
                file_data = uploaded_file.getvalue()
                content = None
                metadata = {
                    'source': new_filename,
                    'original_filename': uploaded_file.name,
                    'type': 'attachment',
                    'file_type': file_extension.strip('.'),
                    'timestamp': timestamp,
                    'processed_date': datetime.datetime.now().isoformat()
                }

                # Process by file type
                if file_extension in ['.png', '.jpg', '.jpeg']:
                    print("\nProcessing Image File:")
                    from PIL import Image
                    import io
                    
                    img = Image.open(io.BytesIO(file_data))
                    print(f"- Image size: {img.size}")
                    print(f"- Image mode: {img.mode}")
                    
                    #Process with Vision API if available
                    vision_analysis = self.process_image(file_data)
                    content = f"Image Analysis:\n{vision_analysis}" if vision_analysis else "No image analysis available"
                    
                    metadata.update({
                        'width': img.size[0],
                        'height': img.size[1],
                        'mode': img.mode,
                        'format': img.format,
                        'size': len(file_data)
                    })
                    
                elif file_extension in ['.docx', '.doc']:
                    print("\nProcessing Word Document:")
                    from docx import Document as DocxDocument
                    import io
                    
                    doc_bytes = io.BytesIO(file_data)
                    doc = DocxDocument(doc_bytes)
                    
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text)
                    
                    tables = []
                    for table in doc.tables:
                        table_data = []
                        for row in table.rows:
                            row_data = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_data.append(cell.text.strip())
                            if row_data:
                                table_data.append(" | ".join(row_data))
                        if table_data:
                            tables.append("\n".join(table_data))
                    
                    content = "\n\n".join(paragraphs)
                    if tables:
                        content += "\n\nTables:\n" + "\n\n".join(tables)
                        
                    metadata.update({
                        'paragraphs': len(paragraphs),
                        'tables': len(doc.tables),
                        'size': len(file_data)
                    })
                    
                elif file_extension in ['.xlsx', '.xls']:
                    print("\nProcessing Excel File:")
                    import pandas as pd
                    import io
                    
                    df = pd.read_excel(io.BytesIO(file_data))
                    content = df.to_string()
                    
                    metadata.update({
                        'rows': len(df),
                        'columns': len(df.columns),
                        'size': len(file_data)
                    })
                    
                elif file_extension == '.pdf':
                    print("\nProcessing PDF File:")
                    import PyPDF2
                    import io
                    
                    pdf_bytes = io.BytesIO(file_data)
                    pdf = PyPDF2.PdfReader(pdf_bytes)
                    content = ""
                    for page in pdf.pages:
                        content += page.extract_text() + "\n"
                        
                    metadata.update({
                        'pages': len(pdf.pages),
                        'size': len(file_data)
                    })
                    
                elif file_extension == '.ipynb':
                    print("\nProcessing Jupyter Notebook:")
                    import json
                    
                    notebook_content = []
                    nb_data = json.loads(file_data.decode('utf-8'))
                    
                    for cell in nb_data['cells']:
                        if cell['cell_type'] in ['markdown', 'code']:
                            if 'source' in cell and cell['source']:
                                if isinstance(cell['source'], list):
                                    notebook_content.append('\n'.join(cell['source']))
                                else:
                                    notebook_content.append(cell['source'])
                    
                    content = '\n\n'.join(notebook_content)
                    
                    metadata.update({
                        'cells': len(nb_data['cells']),
                        'size': len(file_data)
                    })
                    
                else:
                    print("\nProcessing as Text File:")
                    # For other types of files (code, text, etc.)
                    try:
                        content = file_data.decode('utf-8')
                    except UnicodeDecodeError:
                        content = file_data.decode('latin-1')
                        
                    metadata.update({
                        'lines': len(content.splitlines()),
                        'size': len(file_data)
                    })

                if not content:
                    print("No content could be extracted")
                    return {
                        'success': False,
                        'status': 'error',
                        'message': 'No content could be extracted from file'
                    }

                # Calculate tokens
                tokens = self.token_counter.get_tokens(new_filename, content)
                print(f"\nToken Analysis:")
                print(f"- Content length: {len(content)} characters")
                print(f"- Calculated tokens: {tokens}")

                # Update metadata with token information
                metadata.update({
                    'content_length': len(content),
                    'token_count': tokens
                })

                # Create document
                from langchain.schema import Document
                doc = Document(
                    page_content=content,
                    metadata=metadata
                )
                display_text = self.get_file_display(new_filename, tokens)
                self.st.success(display_text, icon="✅")
                # Update session_state
                if 'utils_docs' not in self.st.session_state:
                    self.st.session_state.utils_docs = []
                self.st.session_state.utils_docs.append(doc)
                
                # Update token counters
                if 'file_token_count' not in self.st.session_state:
                    self.st.session_state.file_token_count = {}
                self.st.session_state.file_token_count[new_filename] = tokens
                
                # Update file content
                if 'file_contents' not in self.st.session_state:
                    self.st.session_state.file_contents = {}
                self.st.session_state.file_contents[new_filename] = {
                    'content': content,
                    'metadata': metadata,
                    'timestamp': timestamp
                }

                # Add to embedder if exists
                if hasattr(self.st.session_state, 'embedder'):
                    self.st.session_state.embedder.add_document(doc)
                    print("Added to embedder successfully")

                print("\n=== Processing Summary ===")
                print(f"- Filename: {new_filename}")
                print(f"- Type: {file_extension}")
                print(f"- Size: {metadata['size']} bytes")
                print(f"- Tokens: {tokens}")
                print(f"- Content Length: {len(content)} characters")

                return {
                    'success': True,
                    'status': 'processed',
                    'message': f'Successfully processed {uploaded_file.name}',
                    'filename': new_filename,
                    'tokens': tokens,
                    'metadata': metadata
                }

            except Exception as e:
                print("\n=== Processing Error ===")
                print(f"Error type: {type(e).__name__}")
                print(f"Error message: {str(e)}")
                print("Traceback:")
                import traceback
                traceback.print_exc()
                return {
                    'success': False,
                    'status': 'error',
                    'message': f'Error processing file: {str(e)}'
                }

        except Exception as e:
            print(f"\n=== Critical Error ===")
            print(f"Error type: {type(e).__name__}")
            print(f"Error message: {str(e)}")
            print("Traceback:")
            import traceback
            traceback.print_exc()
            return {
                'success': False,
                'status': 'error',
                'message': f'Critical error: {str(e)}'
            }

    def get_token_breakdown(self):
        """Returns a breakdown of token usage by resource type"""
        breakdown = {
            'jira': 0,
            'confluence': 0,
            'files': 0,
            'github': 0
        }
        
        # Count tokens by document type
        for doc in self.docs:
            doc_type = doc.metadata.get('type', 'files')
            content = doc.page_content
            tokens = self.count_tokens(content)  # Using the unified method
            breakdown[doc_type] += tokens
        
        return breakdown
    
    def count_total_tokens(self, text):
        """Count tokens in text"""
        # Simple approximation: 1 token ≈ 4 characters
        return len(text) // 4 if text else 0
    
    def add_document(self, doc):
        self.docs.append(doc)
        st.session_state.utils_docs = self.docs
     
    def extract_all_files(self):
        """
        Extract and process all files from the cloned repository.
        """
        print("\n=== Starting File Extraction ===")
        try:
            root_dir = "repos"  # Directory where repos are cloned
            total_files = 0
            loaded_files = 0
            allowed_extensions = [
        ".py", ".ipynb", ".md", ".ts", ".yml", ".yaml", ".java", ".xml", ".json",
        ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt", 
        ".png", ".jpg", ".jpeg"
    ]

            if not os.path.exists(root_dir):
                print(f"Directory not found: {root_dir}")
                return False

            for dirpath, dirnames, filenames in os.walk(root_dir):
                # Filter unwanted directories
                dirnames[:] = [d for d in dirnames if not d.startswith('.') and d not in ['node_modules', 'venv', '__pycache__']]
                
                for filename in filenames:
                    file_extension = os.path.splitext(filename)[1]
                    if file_extension.lower() in allowed_extensions:
                        total_files += 1
                        file_path = os.path.join(dirpath, filename)
                        print(f"\nProcessing file: {file_path}")
                        
                        try:
                            # Process README.md with priority
                            if filename.lower() == "readme.md":
                                print("Found README.md - Processing with priority")
                                with open(file_path, 'r', encoding='utf-8') as f:
                                    content = f.read()
                                    if content.strip():
                                        self.docs.append(Document(
                                            page_content=content,
                                            metadata={"source": file_path, "type": "github", "priority": "high"}
                                        ))
                                        loaded_files += 1
                                continue

                           # Process notebooks (.ipynb)
                            if file_extension.lower() == '.ipynb':
                                try:
                                    with open(file_path, 'r', encoding='utf-8') as f:
                                        notebook = json.load(f)
                                        notebook_content = []
                                        for cell in notebook['cells']:
                                            if cell['cell_type'] in ['markdown', 'code']:
                                                if 'source' in cell and cell['source']:
                                                    if isinstance(cell['source'], list):
                                                        notebook_content.append('\n'.join(cell['source']))
                                                    else:
                                                        notebook_content.append(cell['source'])
                                        
                                        if notebook_content:
                                            content = '\n\n'.join(notebook_content)
                                            self.docs.append(Document(
                                                page_content=content,
                                                metadata={"source": file_path, "type": "github", "file_type": "notebook"}
                                            ))
                                            loaded_files += 1
                                    continue
                                except Exception as e:
                                    print(f"Error processing notebook {filename}: {e}")
                                    continue

                           # Process other files
                            with open(file_path, 'r', encoding='utf-8') as f:
                                content = f.read()
                                if content.strip():
                                    self.docs.append(Document(
                                        page_content=content,
                                        metadata={"source": file_path, "type": "github"}
                                    ))
                                    loaded_files += 1
                                    print(f"Successfully loaded {filename}")

                        except Exception as e:
                            print(f"Error processing {filename}: {str(e)}")
                            continue

            print(f"\n=== Extraction Summary ===")
            print(f"Total files found: {total_files}")
            print(f"Successfully loaded files: {loaded_files}")
            print(f"Total documents in memory: {len(self.docs)}")

            return True if loaded_files > 0 else False

        except Exception as e:
            print(f"Critical error in extract_all_files: {str(e)}")
            print(traceback.format_exc())
            return False
    
    def add_to_queue(self, value):
        """
        Adds a value to the chat queue.
        """
        if self.MyQueue.full():
            self.MyQueue.get()
        self.MyQueue.put(value)
       
    def initialize_embedder(self, st) -> Embedder:
        """Create and initialize an Embedder instance."""
        print("Inicializando Embedder en utils...")
        embedder = Embedder(st)
        print("Embedder inicializado con éxito.")
        return embedder

    def load_confluence_page(self, page_url):
        """Load a Confluence page with token counting and enhanced error handling."""
        try:
            print("\n=== Starting Confluence Page Load ===")
            print(f"URL: {page_url}")
            
            #Extract Page ID
            try:
                page_id = None
                if '/pages/' in page_url:
                    parts = page_url.split('/')
                    pages_index = parts.index('pages')
                    page_id = parts[pages_index + 1]
                    print(f"Extracted page ID from URL path: {page_id}")
                else:
                    import re
                    numbers = re.findall(r'/(\d+)/?', page_url)
                    if numbers:
                        page_id = numbers[-1]
                        print(f"Extracted page ID from URL numbers: {page_id}")
                        
                if not page_id:
                    raise ValueError("Could not extract page ID from URL")
                    
            except Exception as e:
                print(f"Error extracting page ID: {str(e)}")
                return False

            # Verify credentials
            api_token = os.getenv('CONFLUENCE_API_KEY')
            email = os.getenv('CONFLUENCE_EMAIL')
            
            if not api_token or not email:
                print("Missing Confluence credentials")
                return False
                
            # Build API URL
            api_url = f"https://theseus-group.atlassian.net/wiki/rest/api/content/{page_id}?expand=body.storage,version,space,ancestors,descendants.page"
            print(f"API URL: {api_url}")

            # Make request to the API
            try:
                response = requests.get(
                    api_url,
                    headers={
                        "Accept": "application/json",
                        "Content-Type": "application/json"
                    },
                    auth=(email, api_token)
                )
                
                print(f"API Response Status: {response.status_code}")
                
                if response.status_code != 200:
                    print(f"API Error: {response.text}")
                    return False
                    
            except requests.exceptions.RequestException as e:
                print(f"Request failed: {str(e)}")
                return False

            # Process response
            try:
                page = response.json()
                
                # Create formatted content
                formatted_content = f"""
    Title: {page['title']}
    Space: {page.get('space', {}).get('name', 'Unknown Space')}

    Content:
    {page['body']['storage']['value']}

    Metadata:
    - Last Updated: {page.get('version', {}).get('when', 'Unknown')}
    - Author: {page.get('version', {}).get('by', {}).get('displayName', 'Unknown')}
    - Version: {page.get('version', {}).get('number', 'Unknown')}
    - URL: {page_url}
    - Page ID: {page_id}

    Additional Information:
    - Space Key: {page.get('space', {}).get('key', 'Unknown')}
    - Created: {page.get('history', {}).get('createdDate', 'Unknown')}
    - Last Modified: {page.get('version', {}).get('when', 'Unknown')}
    """
                
                # Counting tokens
                token_count = self.token_counter.get_tokens(page_url, formatted_content)
                st.success(f"Loaded Confluence page: {page['title']} ({token_count:,} tokens)")
                print(f"\nToken Analysis:")
                print(f"- Page tokens: {token_count}")
                
                # Check token limit
                total_tokens = self.get_total_tokens()
                print(f"- Total tokens after addition: {total_tokens}")
                
                if total_tokens > 16000:
                    print(f"WARNING: Token limit exceeded ({total_tokens}/16000)")
                    self.st.warning(f"⚠️ Adding this page would exceed the token limit ({total_tokens}/16000)")
                    return False

                # Create rich metadata
                metadata = {
                    'source': f"confluence-pages/{page['title']}.md",
                    'title': page['title'],
                    'type': 'confluence',
                    'space': page.get('space', {}).get('name', 'Unknown Space'),
                    'space_key': page.get('space', {}).get('key', 'Unknown'),
                    'url': page_url,
                    'page_id': page_id,
                    'version': page.get('version', {}).get('number', 'Unknown'),
                    'author': page.get('version', {}).get('by', {}).get('displayName', 'Unknown'),
                    'created_date': page.get('history', {}).get('createdDate', 'Unknown'),
                    'last_modified': page.get('version', {}).get('when', 'Unknown'),
                    'content_length': len(formatted_content),
                    'token_count': token_count,
                    'total_tokens': total_tokens,
                    'timestamp': datetime.datetime.now().isoformat()
                }

                # Create document
                document = Document(
                    page_content=formatted_content,
                    metadata=metadata
                )
                
                print("\nDocument created successfully")
                print(f"Metadata: {json.dumps(metadata, indent=2)}")

                # Save to session_state
                if 'confluence_docs' not in self.st.session_state:
                    self.st.session_state.confluence_docs = []
                self.st.session_state.confluence_docs.append(document)
                print(f"Added to session_state.confluence_docs (Total: {len(self.st.session_state.confluence_docs)})")

                # Save to local list
                if not hasattr(self, 'confluence_docs'):
                    self.confluence_docs = []
                self.confluence_docs.append(document)
                print(f"Added to local confluence_docs (Total: {len(self.confluence_docs)})")

                # Add to embedder if exists
                if hasattr(self.st.session_state, 'embedder'):
                    self.st.session_state.embedder.add_document(document)
                    print("Added to embedder successfully")

                # Update statistics
                print("\n=== Page Processing Summary ===")
                print(f"- Title: {page['title']}")
                print(f"- Space: {page.get('space', {}).get('name', 'Unknown Space')}")
                print(f"- Content Length: {len(formatted_content)} characters")
                print(f"- Token Count: {token_count}")
                print(f"- Total System Tokens: {total_tokens}")

                return True

            except Exception as e:
                print(f"\n=== Error Processing Page Content ===")
                print(f"Error type: {type(e).__name__}")
                print(f"Error message: {str(e)}")
                print("Traceback:")
                traceback.print_exc()
                return False

        except Exception as e:
            print(f"\n=== Critical Error in load_confluence_page ===")
            print(f"Error type: {type(e).__name__}")
            print(f"Error message: {str(e)}")
            print("Traceback:")
            traceback.print_exc()
            return False

    def load_jira_ticket(self, jira_url):
        """Load and process a Jira ticket with token counting and enhanced error handling."""
        try:
            print("\n=== Starting Jira Ticket Load ===")
            print(f"URL: {jira_url}")
            
            # Extract Issue Key
            try:
                if 'selectedIssue=' in jira_url:
                    issue_key = jira_url.split('selectedIssue=')[-1]
                    print(f"Extracted issue key from selectedIssue: {issue_key}")
                else:
                    issue_key = jira_url.split('/')[-1]
                    print(f"Extracted issue key from URL path: {issue_key}")
                    
                if not issue_key:
                    raise ValueError("Could not extract issue key from URL")
                    
            except Exception as e:
                print(f"Error extracting issue key: {str(e)}")
                return False

            # Verify credentials
            api_token = os.getenv('CONFLUENCE_API_KEY')
            email = os.getenv('CONFLUENCE_EMAIL')
            
            if not api_token or not email:
                print("Missing Jira credentials")
                return False
                
            # Verify credentials
            api_url = f"https://theseus-group.atlassian.net/rest/api/3/issue/{issue_key}"
            print(f"API URL: {api_url}")

            # Headers for rich information
            headers = {
                "Accept": "application/json",
                "Content-Type": "application/json"
            }

            # Make request to the API
            try:
                response = requests.get(
                    api_url,
                    headers=headers,
                    auth=(email, api_token)
                )
                
                print(f"API Response Status: {response.status_code}")
                
                if response.status_code != 200:
                    print(f"API Error: {response.text}")
                    return False
                    
            except requests.exceptions.RequestException as e:
                print(f"Request failed: {str(e)}")
                return False

            try:
                issue_data = response.json()
                if not isinstance(issue_data, dict):
                    print("Error: Invalid response format from Jira API")
                    return False

                fields = issue_data.get('fields', {})
                if fields is None:
                    print("Error: Fields is None in Jira response")
                    fields = {}  
                    
                if not issue_data.get('key'):
                    print("Error: Missing issue key in response")
                    return False
                    
                if not fields.get('summary'):
                    print("Warning: Missing summary in fields")
                    
                if fields.get('status') is None:
                    fields['status'] = {}
                if fields.get('priority') is None:
                    fields['priority'] = {}
                if fields.get('issuetype') is None:
                    fields['issuetype'] = {}
                if fields.get('assignee') is None:
                    fields['assignee'] = {}
                if fields.get('reporter') is None:
                    fields['reporter'] = {}
                                
                # Extract comments if they exist
                comments = []
                if fields.get('comment'):
                    for comment in fields['comment'].get('comments', []):
                        comments.append({
                            'author': comment.get('author', {}).get('displayName', 'Unknown'),
                            'created': comment.get('created', 'Unknown'),
                            'body': comment.get('body', 'No content')
                        })

                # Extract attachments if they exist
                attachments = []
                if fields.get('attachment'):
                    for attachment in fields['attachment']:
                        attachments.append({
                            'filename': attachment.get('filename', 'Unknown'),
                            'size': attachment.get('size', 0),
                            'created': attachment.get('created', 'Unknown')
                        })

                # Create formatted content with safe gets
                formatted_content = f"""
    Issue Key: {issue_data.get('key', 'Unknown')}
    Summary: {fields.get('summary', 'No summary')}

    Status: {fields.get('status', {}).get('name', 'Unknown')}
    Priority: {fields.get('priority', {}).get('name', 'Unknown')}
    Type: {fields.get('issuetype', {}).get('name', 'Unknown')}

    Description:
    {fields.get('description', 'No description')}

    Timeline:
    - Created: {fields.get('created', 'Unknown')}
    - Updated: {fields.get('updated', 'Unknown')}
    - Due Date: {fields.get('duedate', 'Not set')}

    Assignee: {fields.get('assignee', {}).get('displayName', 'Unassigned')}
    Reporter: {fields.get('reporter', {}).get('displayName', 'Unknown')}

    Comments:
    """
                # Add comments to content
                if comments:
                    for comment in comments:
                        formatted_content += f"""
    - Author: {comment['author']}
    Date: {comment['created']}
    Content: {comment['body']}
    """
                else:
                    formatted_content += "No comments yet\n"

                # Add attachments to content
                formatted_content += "\nAttachments:\n"
                if attachments:
                    for attachment in attachments:
                        formatted_content += f"- {attachment['filename']} ({attachment['size']} bytes, added: {attachment['created']})\n"
                else:
                    formatted_content += "No attachments\n"
                
                # Counting tokens
                token_count = self.token_counter.get_tokens(jira_url, formatted_content)
                st.success(f"Loaded Jira ticket: {issue_data.get('key', 'Unknown')} ({token_count:,} tokens)")
                print(f"\nToken Analysis:")
                print(f"- Issue tokens: {token_count}")
                
                # Check token limit
                total_tokens = self.get_total_tokens()
                print(f"- Total tokens after addition: {total_tokens}")
                
                if total_tokens > 16000:
                    print(f"WARNING: Token limit exceeded ({total_tokens}/16000)")
                    self.st.warning(f"⚠️ Adding this issue would exceed the token limit ({total_tokens}/16000)")
                    return False

                # Create rich metadata with safe gets
                metadata = {
                    'source': f"jira-issues/{issue_data.get('key', 'unknown')}.md",
                    'title': f"{issue_data.get('key', 'Unknown')}: {fields.get('summary', 'No summary')}",
                    'type': 'jira',
                    'url': jira_url,
                    'issue_key': issue_data.get('key', 'Unknown'),
                    'status': fields.get('status', {}).get('name', 'Unknown'),
                    'priority': fields.get('priority', {}).get('name', 'Unknown'),
                    'issue_type': fields.get('issuetype', {}).get('name', 'Unknown'),
                    'created_date': fields.get('created', 'Unknown'),
                    'updated_date': fields.get('updated', 'Unknown'),
                    'due_date': fields.get('duedate', 'Not set'),
                    'assignee': fields.get('assignee', {}).get('displayName', 'Unassigned'),
                    'reporter': fields.get('reporter', {}).get('displayName', 'Unknown'),
                    'comment_count': len(comments),
                    'attachment_count': len(attachments),
                    'content_length': len(formatted_content),
                    'token_count': token_count,
                    'total_tokens': total_tokens,
                    'timestamp': datetime.datetime.now().isoformat()
                }

                # Create document
                document = Document(
                    page_content=formatted_content,
                    metadata=metadata
                )
                
                print("\nDocument created successfully")
                print(f"Metadata: {json.dumps(metadata, indent=2)}")

                # Save to session_state
                if 'jira_docs' not in self.st.session_state:
                    self.st.session_state.jira_docs = []
                self.st.session_state.jira_docs.append(document)
                print(f"Added to session_state.jira_docs (Total: {len(self.st.session_state.jira_docs)})")

                # Save to local list
                if not hasattr(self, 'jira_docs'):
                    self.jira_docs = []
                self.jira_docs.append(document)
                print(f"Added to local jira_docs (Total: {len(self.jira_docs)})")

                # Add to embedder if exists
                if hasattr(self.st.session_state, 'embedder'):
                    self.st.session_state.embedder.add_document(document)
                    print("Added to embedder successfully")

                # Update statistics
                print("\n=== Issue Processing Summary ===")
                print(f"- Key: {issue_data.get('key', 'Unknown')}")
                print(f"- Summary: {fields.get('summary', 'No summary')}")
                print(f"- Status: {fields.get('status', {}).get('name', 'Unknown')}")
                print(f"- Comments: {len(comments)}")
                print(f"- Attachments: {len(attachments)}")
                print(f"- Content Length: {len(formatted_content)} characters")
                print(f"- Token Count: {token_count}")
                print(f"- Total System Tokens: {total_tokens}")

                # Show successful toast
                self.st.toast(f"✅ Loaded Jira issue: {issue_data.get('key', 'Unknown')}")
                return True

            except Exception as e:
                print(f"\n=== Error Processing Issue Content ===")
                print(f"Error type: {type(e).__name__}")
                print(f"Error message: {str(e)}")
                print("Traceback:")
                traceback.print_exc()
                self.st.toast(f"❌ Error loading issue: {str(e)}")
                return False

        except Exception as e:
            print(f"\n=== Critical Error in load_jira_ticket ===")
            print(f"Error type: {type(e).__name__}")
            print(f"Error message: {str(e)}")
            print("Traceback:")
            traceback.print_exc()
            return False

    def process_document_content(self, file_obj, file_type):
        """Processes the content of documents in a type-specific manner"""
        try:
            print(f"\n=== Processing Document Content ===")
            print(f"File type: {file_type}")
            print(f"Processing: {file_obj.name}")
            
            content = None
            metadata = {}
            
            if file_type == 'docx':
                try:
                    from docx import Document
                    doc = Document(file_obj)
                    
                    # Extract text from paragraphs
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text)
                    
                    # Extract text from tables
                    tables = []
                    for table in doc.tables:
                        table_data = []
                        for row in table.rows:
                            row_data = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_data.append(cell.text.strip())
                            if row_data:
                                table_data.append(" | ".join(row_data))
                        if table_data:
                            tables.append("\n".join(table_data))
                    
                    # Combine all content
                    content = "\n\n".join(paragraphs)
                    if tables:
                        content += "\n\nTables:\n" + "\n\n".join(tables)
                    
                    metadata = {
                        'paragraphs': len(paragraphs),
                        'tables': len(doc.tables),
                        'has_content': bool(content)
                    }
                    
                    print(f"Extracted {len(paragraphs)} paragraphs and {len(doc.tables)} tables")
                    
                except Exception as e:
                    print(f"Error processing DOCX: {str(e)}")
                    raise
                    
            # Add the processed document to the embedder
            if content:
                try:
                    doc = Document(
                        page_content=content,
                        metadata={
                            'source': file_obj.name,
                            'type': 'document',
                            'file_type': file_type,
                            **metadata
                        }
                    )
                    
                    # Add to embedder
                    if hasattr(st.session_state, 'embedder'):
                        st.session_state.embedder.add_document(doc)
                        print("Document added to embedder successfully")
                    
                    # Add to utils_docs
                    if 'utils_docs' not in st.session_state:
                        st.session_state.utils_docs = []
                    st.session_state.utils_docs.append(doc)
                    
                    return {
                        'success': True,
                        'content': content,
                        'metadata': metadata
                    }
                    
                except Exception as e:
                    print(f"Error adding document to embedder: {str(e)}")
                    raise
            else:
                print("No content extracted from document")
                return {
                    'success': False,
                    'error': 'No content could be extracted'
                }
                
        except Exception as e:
            print(f"Error in process_document_content: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }


    def retrieve_results(self, query, image=None, selected_files=None):
        try:
            print("\n=== Processing Query ===")
            print(f"Query: {query}")
            
            # Collect all available documents
            all_docs = []
            
            # Add utils_docs documents (includes uploaded files)
            if 'utils_docs' in self.st.session_state:
                print(f"Adding {len(self.st.session_state.utils_docs)} utils_docs")
                all_docs.extend(self.st.session_state.utils_docs)
                
            # Adding Jira Documents
            if 'jira_docs' in self.st.session_state:
                print(f"Adding {len(self.st.session_state.jira_docs)} jira_docs")
                all_docs.extend(self.st.session_state.jira_docs)
                
            # Adding Confluence Documents
            if 'confluence_docs' in self.st.session_state:
                print(f"Adding {len(self.st.session_state.confluence_docs)} confluence_docs")
                all_docs.extend(self.st.session_state.confluence_docs)

            print(f"Total documents available: {len(all_docs)}")
            
            # Building the context
            context = ""
            for doc in all_docs:
                source = doc.metadata.get('source', 'Unknown')
                content = doc.page_content
                context += f"\nSource: {source}\nContent:\n{content}\n"
                print(f"Added content from: {source}")

            # Preparing the prompt
            system_prompt = """You are a helpful AI assistant that analyzes documents and answers questions about them. 
            Base your answers only on the provided context and be specific about your sources."""

            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Context:\n{context}\n\nQuestion: {query}"}
            ]

            # Calling the API
            client = OpenAI()
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=messages,
                temperature=0.7,
                max_tokens=2000
            )

            return response.choices[0].message.content

        except Exception as e:
            print(f"Error in retrieve_results: {str(e)}")
            return f"Error processing query: {str(e)}"

    def process_image(self, image):
        """Process an image using the Vision API."""
        try:
            print("\n=== Starting Image Processing ===")
            print(f"Initial image type: {type(image)}")
            print(f"Image details: {image}")
            
            client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
            print("OpenAI client initialized")
            
            # Verify and convert image to base64
            print("\n=== Converting Image ===")
            try:
                if isinstance(image, bytes):
                    print("Image is in bytes format")
                    image_base64 = base64.b64encode(image).decode('utf-8')
                    print("Successfully converted bytes to base64")
                else:
                    print("Image is not in bytes format, converting from PIL Image")
                    img_byte_arr = io.BytesIO()
                    print("Created BytesIO object")
                    
                    image.save(img_byte_arr, format='PNG')
                    print("Saved image to BytesIO")
                    
                    image_base64 = base64.b64encode(img_byte_arr.getvalue()).decode('utf-8')
                    print("Successfully converted to base64")
                    
                    img_byte_arr.close()
                    print("Closed BytesIO object")
                
                print(f"Base64 string length: {len(image_base64)}")
                print(f"Base64 string preview: {image_base64[:100]}...")
            except Exception as e:
                print(f"Error in image conversion: {str(e)}")
                raise

            print("\n=== Making API Call ===")
            print("Preparing API request...")
            try:
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": "Analyze this image in detail and describe what you see."
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/png;base64,{image_base64}"
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=1000
                )
                print("API call completed successfully")
                print(f"Response structure: {response}")
                
                content = response.choices[0].message.content
                print(f"Generated content length: {len(content)}")
                print(f"Content preview: {content[:200]}")
                
                return content
                
            except Exception as e:
                print("\n=== API Call Error ===")
                print(f"Error type: {type(e).__name__}")
                print(f"Error message: {str(e)}")
                print("API request details:")
                print(f"- Model: gpt-4o")
                print(f"- Max tokens: 1000")
                print(f"- API key length: {len(os.getenv('OPENAI_API_KEY'))}")
                raise

        except Exception as e:
            print(f"\n=== Error in process_image ===")
            print(f"Error type: {type(e).__name__}")
            print(f"Error message: {str(e)}")
            print("Traceback:")
            traceback.print_exc()
            return None


    def clone_and_process_github(self, git_url):
        """
        Clone and process a GitHub repository with detailed logs.
        """
        print(f"\n=== Processing GitHub Repository ===")
        print(f"URL: {git_url}")
        
        try:
            # Validate URL
            if not git_url or 'github.com' not in git_url:
                print("Invalid GitHub URL")
                return False

            # Prepare path
            repo_name = git_url.split('/')[-1].split('.')[0]
            clone_path = f"repos/{repo_name}"
            print(f"Repository will be cloned to: {clone_path}")
            
            # Clone or update repo
            if os.path.exists(clone_path):
                print(f"Repository exists at {clone_path}, updating...")
                repo = git.Repo(clone_path)
                repo.remotes.origin.pull()
                print("Repository updated successfully")
            else:
                print(f"Cloning new repository to {clone_path}")
                repo = git.Repo.clone_from(git_url, clone_path)
                print("Repository cloned successfully")
            
            # Process files
            print("\n=== Processing Files ===")
            allowed_extensions = [".py", ".ipynb", ".md", ".ts", ".yml", ".java", ".xml", ".json"]
            files_processed = 0
            
            for dirpath, dirnames, filenames in os.walk(clone_path):
                #Filter hidden directories
                dirnames[:] = [d for d in dirnames if not d.startswith('.') and d not in ['node_modules', 'venv', '__pycache__']]
                
                for filename in filenames:
                    file_path = os.path.join(dirpath, filename)
                    if any(filename.endswith(ext) for ext in allowed_extensions):
                        print(f"\nProcessing: {filename}")
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                content = f.read()
                                
                            doc = Document(
                                page_content=content,
                                metadata={
                                    "source": file_path,
                                    "type": "github",
                                    "filename": filename,
                                    "repository": repo_name
                                }
                            )
                            self.docs.append(doc)
                            files_processed += 1
                            print(f"Successfully processed: {filename}")
                            
                        except Exception as e:
                            print(f"Error processing file {filename}: {str(e)}")
                            continue

            print(f"\n=== Repository Processing Summary ===")
            print(f"Total files processed: {files_processed}")
            print(f"Total documents in memory: {len(self.docs)}")
            
            return True if files_processed > 0 else False

        except git.exc.GitCommandError as e:
            print(f"\n=== Git Command Error ===")
            print(f"Command that failed: {e.command}")
            print(f"Status code: {e.status}")
            print(f"Error message: {e.stderr}")
            return False
            
        except Exception as e:
            print(f"\n=== Unexpected Error ===")
            print(f"Error type: {type(e).__name__}")
            print(f"Error message: {str(e)}")
            import traceback
            print(f"Traceback: {traceback.format_exc()}")
            return False 

    def load_file_content(self, file_path):
        """Load and process file content with enhanced token counting and error handling."""
        try:
            if not os.path.exists(file_path):
                return False
                
            base_path = os.path.normpath(file_path)
            _, ext = os.path.splitext(base_path)
            
            allowed_extensions = {
                'code': ['.py', '.java', '.ts', '.xml', '.json', '.yml', '.yaml', '.sh'],
                'document': ['.docx', '.doc', '.xlsx', '.xls', '.pdf'],
                'text': ['.txt', '.md', '.requirements'],
                'notebook': ['.ipynb'],
                'image': ['.png', '.jpg', '.jpeg']
            }

            file_type = None
            for type_name, extensions in allowed_extensions.items():
                if ext.lower() in extensions:
                    file_type = type_name
                    break

            if file_type is None:
                file_type = 'unknown'
                self.st.info(f"Note: File type '{ext}' is not explicitly supported but will be processed as text")

            for doc in self.docs:
                if doc.metadata.get('source') == base_path:
                    return True

            try:
                content = None
                
                file_stat = os.stat(base_path)
                metadata = {
                    'source': base_path,
                    'title': os.path.basename(base_path),
                    'extension': ext,
                    'file_type': file_type,
                    'last_modified': datetime.datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
                    'file_size': file_stat.st_size,
                    'processed_timestamp': datetime.datetime.now().isoformat()
                }

                if file_type == 'document':
                    if ext.lower() in ['.docx', '.doc']:
                        from docx import Document as DocxDocument
                        doc = DocxDocument(base_path)
                        
                        paragraphs = []
                        for para in doc.paragraphs:
                            if para.text.strip():
                                paragraphs.append(para.text)
                        
                        tables = []
                        for table in doc.tables:
                            table_data = []
                            for row in table.rows:
                                row_data = []
                                for cell in row.cells:
                                    if cell.text.strip():
                                        row_data.append(cell.text.strip())
                                if row_data:
                                    table_data.append(" | ".join(row_data))
                            if table_data:
                                tables.append("\n".join(table_data))
                        
                        content = "\n\n".join(paragraphs)
                        if tables:
                            content += "\n\nTables:\n" + "\n\n".join(tables)
                        
                        metadata.update({
                            'paragraphs': len(paragraphs),
                            'tables': len(doc.tables)
                        })

                    elif ext.lower() in ['.xlsx', '.xls']:
                        import pandas as pd
                        df = pd.read_excel(base_path)
                        content = df.to_string()
                        metadata.update({
                            'rows': len(df),
                            'columns': len(df.columns)
                        })

                    elif ext.lower() == '.pdf':
                        import PyPDF2
                        with open(base_path, 'rb') as file:
                            reader = PyPDF2.PdfReader(file)
                            content = ''
                            for page in reader.pages:
                                content += page.extract_text() + '\n'
                            metadata['pages'] = len(reader.pages)

                elif file_type == 'image':
                    from PIL import Image
                    image = Image.open(base_path)
                    metadata.update({
                        'dimensions': f"{image.size[0]}x{image.size[1]}",
                        'image_mode': image.mode
                    })
                    
                    img_byte_arr = io.BytesIO()
                    image.save(img_byte_arr, format='PNG')
                    img_byte_arr = img_byte_arr.getvalue()
                    
                    vision_analysis = self.process_image(img_byte_arr)
                    content = f"Image Analysis:\n{vision_analysis}" if vision_analysis else "No image analysis available"

                elif file_type == 'notebook':
                    notebook_content = []
                    with open(base_path, 'r', encoding='utf-8') as f:
                        notebook = json.load(f)
                        for cell in notebook['cells']:
                            if cell['cell_type'] in ['markdown', 'code']:
                                if 'source' in cell and cell['source']:
                                    if isinstance(cell['source'], list):
                                        notebook_content.append('\n'.join(cell['source']))
                                    else:
                                        notebook_content.append(cell['source'])
                        content = '\n\n'.join(notebook_content)
                        metadata['cells'] = len(notebook['cells'])

                else:
                    try:
                        with open(base_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                    except UnicodeDecodeError:
                        with open(base_path, 'r', encoding='latin-1') as f:
                            content = f.read()
                    metadata['lines'] = content.count('\n') + 1

                if not content:
                    return False

                tokens = self.process_file_tokens(base_path)

                metadata.update({
                    'content_length': len(content),
                    'token_count': tokens,
                    'tokens_per_character': tokens / len(content) if len(content) > 0 else 0
                })

                total_tokens = self.get_total_tokens()
                if total_tokens > 14000:
                    self.st.warning(f"⚠️ High token usage: {total_tokens:,}/16,000")

                document = Document(
                    page_content=content,
                    metadata=metadata
                )

                if hasattr(self.st.session_state, 'embedder'):
                    self.st.session_state.embedder.add_document(document)

                if 'utils_docs' not in self.st.session_state:
                    self.st.session_state.utils_docs = []
                self.st.session_state.utils_docs.append(document)

                self.docs.append(document)

                self.token_counts['files'][base_path] = tokens
                if 'file_token_count' not in self.st.session_state:
                    self.st.session_state.file_token_count = {}
                self.st.session_state.file_token_count[base_path] = tokens

                if 'file_contents' not in self.st.session_state:
                    self.st.session_state.file_contents = {}
                self.st.session_state.file_contents[base_path] = content

                self.st.toast(
                    self.get_file_display(base_path, tokens),
                    icon="✅"
                )

                return True

            except Exception as e:
                self.st.toast(f"❌ Error processing {os.path.basename(base_path)}: {str(e)}", icon="❌")
                return False

        except Exception as e:
            return False

    def clone_repo(self, git_url):
            """
            Alias ​​for clone_and_process_github with additional logs.
            """
            print("\n=== Starting GitHub Repository Clone ===")
            print(f"Attempting to clone repository: {git_url}")
            try:
                if not os.path.exists("repos"):
                    print("Creating repos directory...")
                    os.makedirs("repos")
                    
                return self.clone_and_process_github(git_url)
            except Exception as e:
                print(f"Error in clone_repo: {str(e)}")
                return False

    
    def remove_file(self, file_id):
        "Delete file and its related records"""
        try:
            print(f"\n=== Removing file: {file_id} ===")
            
            # Remove from utils_docs
            if 'utils_docs' in self.st.session_state:
                self.st.session_state.utils_docs = [
                    doc for doc in self.st.session_state.utils_docs 
                    if doc.metadata.get('source') != file_id
                ]
                print(f"Removed from utils_docs")
            
            # Remove from file_token_count
            if 'file_token_count' in self.st.session_state:
                if file_id in self.st.session_state.file_token_count:
                    del self.st.session_state.file_token_count[file_id]
                    print(f"Removed from file_token_count")
            
            # Delete from file_contents
            if 'file_contents' in self.st.session_state:
                if file_id in self.st.session_state.file_contents:
                    del self.st.session_state.file_contents[file_id]
                    print(f"Removed from file_contents")
            
            # Remove from selected resources
            if 'selected_resources' in self.st.session_state:
                if 'files' in self.st.session_state.selected_resources:
                    if file_id in self.st.session_state.selected_resources['files']:
                        del self.st.session_state.selected_resources['files'][file_id]
                        print(f"Removed from selected_resources")
            
            # Remove from embedder if exists
            if hasattr(self.st.session_state, 'embedder'):
                self.st.session_state.embedder.docs = [
                    doc for doc in self.st.session_state.embedder.docs 
                    if doc.metadata.get('source') != file_id
                ]
                print(f"Removed from embedder")
            
            print(f"File {file_id} successfully removed from all locations")
            return True
            
        except Exception as e:
            print(f"Error removing file {file_id}: {str(e)}")
            return False

    
    def process_query(self, query, mode="chat"):
        """Process the query according to the mode with token management"""
        try:
            print(f"\n=== Processing Query ===")
            print(f"Mode: {mode}")
            print(f"Query: {query}")
            
            # Check prompt length
            prompt_tokens = self.update_token_count('prompt', 'current_query', query)
            print(f"Query tokens: {prompt_tokens}")
            
            # Get current token status
            current_total = self.get_total_tokens()
            available_tokens = 16000 - current_total
            print(f"Current total tokens: {current_total}")
            print(f"Available tokens: {available_tokens}")
            
            if prompt_tokens > available_tokens:
                print(f"Query too long: {prompt_tokens} tokens needed, only {available_tokens} available")
                return f"⚠️ Query too long. Please reduce length (needs {prompt_tokens} tokens, only {available_tokens} available)"
            
            if mode == "chat":
                print("\n=== Chat Mode Processing ===")
                # Get current state
                print("Checking session state...")
                messages = self.st.session_state.get('messages', [])
                print(f"Current messages in state: {len(messages)}")
                
                image = self.st.session_state.get('image_data')
                if image:
                    print("Image data found in session")
                
                selected_files = self.st.session_state.tree_state.get("checked", [])
                print(f"Selected files: {selected_files}")
                
                # Get available documents
                all_docs = []
                if 'utils_docs' in self.st.session_state:
                    print(f"Adding {len(self.st.session_state.utils_docs)} utils_docs")
                    all_docs.extend(self.st.session_state.utils_docs)
                    
                if 'jira_docs' in self.st.session_state:
                    print(f"Adding {len(self.st.session_state.jira_docs)} jira_docs")
                    all_docs.extend(self.st.session_state.jira_docs)
                    
                if 'confluence_docs' in self.st.session_state:
                    print(f"Adding {len(self.st.session_state.confluence_docs)} confluence_docs")
                    all_docs.extend(self.st.session_state.confluence_docs)
                
                print(f"Total documents available: {len(all_docs)}")
                
                # Build context
                print("\nBuilding context...")
                context = ""
                context_tokens = 0
                
                for doc in all_docs:
                    source = doc.metadata.get('source', 'Unknown')
                    content = doc.page_content
                    doc_context = f"\nSource: {source}\nContent:\n{content}\n"
                    doc_tokens = self.count_tokens(doc_context)
                    
                    if context_tokens + doc_tokens + prompt_tokens > 14000:  
                        print(f"Skipping document {source} to avoid token limit")
                        continue
                        
                    context += doc_context
                    context_tokens += doc_tokens
                    print(f"Added content from: {source} ({doc_tokens} tokens)")
                
                print(f"Total context tokens: {context_tokens}")
                
                # Prepare messages
                print("\nPreparing chat messages...")
                system_prompt = """You are a helpful AI assistant that analyzes documents and answers questions about them. 
                Base your answers only on the provided context and be specific about your sources."""
                
                system_tokens = self.count_tokens(system_prompt)
                total_input_tokens = system_tokens + context_tokens + prompt_tokens
                
                if total_input_tokens > 14000:
                    return "⚠️ Combined input too large. Please reduce the number of selected documents or simplify your query."
                
                messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Context:\n{context}\n\nQuestion: {query}"}
                ]
                
                print(f"Total input tokens: {total_input_tokens}")
                print("Making API call...")
                
                try:
                    client = OpenAI()
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=messages,
                        temperature=0.7,
                        max_tokens=min(2000, 16000 - total_input_tokens)  # Adjust max_tokens dynamically
                    )
                    print("API call successful")
                    
                    # Add messages to session state
                    if 'messages' not in self.st.session_state:
                        self.st.session_state.messages = []
                        
                    self.st.session_state.messages.append({"role": "user", "content": query})
                    response_content = response.choices[0].message.content
                    self.st.session_state.messages.append({"role": "assistant", "content": response_content})
                    
                    # Update token counters
                    response_tokens = self.count_tokens(response_content)
                    self.update_token_count('response', 'last_response', response_content)
                    
                    print(f"Response tokens: {response_tokens}")
                    print(f"Total tokens used: {total_input_tokens + response_tokens}")
                    
                    return response_content
                    
                except Exception as e:
                    print(f"API call error: {str(e)}")
                    return f"Error calling API: {str(e)}"
                
            else:
                # Single Prompt Mode
                print("\n=== Single Prompt Mode ===")
                embedder = self.st.session_state.embedder
                
                # Check available documents
                print(f"Documents in embedder: {len(embedder.docs)}")
                
                # Build context with token counting
                context = ""
                context_tokens = 0
                
                for doc in embedder.docs:
                    source = doc.metadata.get('source', 'Unknown')
                    content = doc.page_content
                    doc_context = f"\nSource: {source}\nContent:\n{content}\n"
                    doc_tokens = self.count_tokens(doc_context)
                    
                    if context_tokens + doc_tokens + prompt_tokens > 14000:
                        print(f"Skipping document {source} to avoid token limit")
                        continue
                        
                    context += doc_context
                    context_tokens += doc_tokens
                    print(f"Added content from: {source} ({doc_tokens} tokens)")

                #Preparing the prompt
                system_prompt = """You are a helpful AI assistant that analyzes documents and answers questions about them. 
                Base your answers only on the provided context and be specific about your sources."""
                
                system_tokens = self.count_tokens(system_prompt)
                total_input_tokens = system_tokens + context_tokens + prompt_tokens
                
                if total_input_tokens > 14000:
                    return "⚠️ Combined input too large. Please reduce the number of selected documents or simplify your query."

                messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Context:\n{context}\n\nQuestion: {query}"}
                ]

                # Calling the API
                try:
                    client = OpenAI()
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=messages,
                        temperature=0.7,
                        max_tokens=min(2000, 16000 - total_input_tokens)
                    )
                    
                    response_content = response.choices[0].message.content
                    response_tokens = self.count_tokens(response_content)
                    
                    print(f"Response tokens: {response_tokens}")
                    print(f"Total tokens used: {total_input_tokens + response_tokens}")
                    
                    return response_content
                    
                except Exception as e:
                    print(f"API call error: {str(e)}")
                    return f"Error calling API: {str(e)}"

        except Exception as e:
            print(f"\n=== Error in process_query ===")
            print(f"Error type: {type(e).__name__}")
            print(f"Error message: {str(e)}")
            print("Traceback:")
            traceback.print_exc()
            return f"Error processing query: {str(e)}"


    def display_file_tokens(self, file_path):
        """Show token count for a file"""
        normalized_path = os.path.normpath(file_path)
        tokens = self.token_counts['files'].get(normalized_path, 0)
        
        if tokens > 0:
            return f"{os.path.basename(file_path)} ({tokens:,} tokens)"
        else:
            # If there are no tokens, try to calculate them
            tokens = self.process_file_tokens(file_path)
            return f"{os.path.basename(file_path)} ({tokens:,} tokens)"

    def analyze_path(self, path):
        """Analyzes a specific route and its contents."""
        try:
            print(f"\n=== Analyzing path: {path} ===")
            if not os.path.exists(path):
                return f"The path {path} does not exist."
            
            if os.path.isfile(path):
                return self.analyze_file(path)
            
            # If it is a directory
            files = []
            subdirs = []
            total_size = 0
            
            for entry in os.scandir(path):
                if entry.is_file():
                    size = entry.stat().st_size
                    total_size += size
                    files.append({
                        'name': entry.name,
                        'size': size,
                        'last_modified': datetime.fromtimestamp(entry.stat().st_mtime).strftime('%Y-%m-%d %H:%M:%S')
                    })
                elif entry.is_dir():
                    subdirs.append(entry.name)
            
            response = f"""
    Route analysis: {path}
   Type: Directory
    Number of files:{len(files)}
    Number of subdirectories: {len(subdirs)}
    Overall size: {total_size / 1024:.2f} KB
    Main files:
    """
            
            for file in sorted(files, key=lambda x: x['size'], reverse=True)[:5]:
                response += f"- {file['name']} ({file['size'] / 1024:.2f} KB, modificado: {file['last_modified']})\n"
            
            if subdirs:
                response += "\nSubdirectorios:\n"
                for subdir in sorted(subdirs):
                    response += f"- {subdir}\n"
            
            return response
            
        except Exception as e:
            print(f"Error analyzing path: {str(e)}")
            return f"Error parsing path: {str(e)}"
        
    def analyze_file(self, file_path):
        """Analyze a specific file."""
        try:
            if not os.path.exists(file_path):
                return f"The file {file_path} does not exists."
            
            stats = os.stat(file_path)
            _, ext = os.path.splitext(file_path)
            
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    lines = content.count('\n') + 1
            except:
                content = None
                lines = "Not readable as text"
            
            response = f"""
    File analysis: {os.path.basename(file_path)}
    Full route: {file_path}
    Type: {ext}
    Size: {stats.st_size / 1024:.2f} KB
    Last Modification: {datetime.fromtimestamp(stats.st_mtime).strftime('%Y-%m-%d %H:%M:%S')}
    Lines: {lines}
    """
            if content and ext.lower() in ['.txt', '.md', '.py', '.json', '.yml', '.yaml']:
                response += "Content:\n```\n"
                response += content[:1000] + ("..." if len(content) > 1000 else "")
                response += "\n```"
            
            return response
            
        except Exception as e:
            print(f"Error analyzing file: {str(e)}")
            return f"Error parsing file: {str(e)}"
        
    def is_file_processed(self, file_path):
        """Check if a file is already processed using normalized path"""
        try:
            if not os.path.exists(file_path):
                return False
                    
            normalized_path = os.path.normpath(file_path)
            existing_paths = set()
            
            if hasattr(self.st.session_state, 'utils_docs'):
                utils_paths = {os.path.normpath(doc.metadata.get('source', '')) 
                            for doc in self.st.session_state.utils_docs}
                existing_paths.update(utils_paths)
            
            local_paths = {os.path.normpath(doc.metadata.get('source', ''))
                        for doc in self.docs}
            existing_paths.update(local_paths)
            
            if 'processed_files' in self.st.session_state:
                processed_paths = {os.path.normpath(path)
                                for path in self.st.session_state.processed_files}
                existing_paths.update(processed_paths)
            
            if 'files' in self.token_counts:
                token_paths = {os.path.normpath(path)
                            for path in self.token_counts['files'].keys()}
                existing_paths.update(token_paths)
            
            existing_paths.discard('')
            existing_paths.discard(None)
            
            return normalized_path in existing_paths
                
        except Exception as e:
            return False


    def get_file_icon(self, file_path):
        try:
            # Special cases for specific directories
            if file_path == 'folder' or (os.path.exists(file_path) and os.path.isdir(file_path)):
                dir_name = os.path.basename(file_path) if file_path != 'folder' else 'folder'
                
                # Specific mapping for special directories
                special_dirs = {
                    '__pycache__': 'python_cache.svg',  
                    'repos': 'repository.svg',         
                    'node_modules': 'node.svg',         
                    'venv': 'python_env.svg',           
                }
                
                icon_name = special_dirs.get(dir_name, 'folder.svg')
            else:
                # Get file extension for regular files
                extension = Path(file_path).suffix.lower().replace('.', '')

                # Mapping of file extensions to icons
                icon_map = {
                    'py': 'file_type_python.svg',
                    'md': 'file_type_markdown.svg',
                    'txt': 'file_type_text.svg',
                    'json': 'file_type_json.svg',
                    'yml': 'file_type_yaml.svg',
                    'yaml': 'file_type_yaml.svg',
                    'html': 'file_type_html.svg',
                    'csv': 'file_type_csv.svg',
                    'docx': 'file_type_docx.svg',
                    'xlsx': 'file_type_xlsx.svg',
                    'pdf': 'file_type_pdf.svg',
                    'png': 'file_type_image.svg',
                    'jpg': 'file_type_image.svg',
                    'jpeg': 'file_type_image.svg',
                    'zip': 'file_type_zip.svg',
                    'rar': 'file_type_rar.svg',
                    'exe': 'file_type_exe.svg',
                    'folder': 'folder.svg',
                    'jira': 'jira.svg',
                    'confluence': 'confluence.svg',
                    'cpython.pyc':'file_type_cython.svg',
                    'js':'file_type_jss.svg',
                    'css':'file_type_css.svg',
                    'log':'file_type_log.svg',
                    'sh':'file_type_shell.svg'
                }
                
                icon_name = icon_map.get(extension, 'default.svg')

            icons_path = Path(__file__).parent / 'icons'
            icon_path = icons_path / icon_name

            if not icon_path.exists():
                return None

            with open(icon_path, 'rb') as f:
                encoded = base64.b64encode(f.read()).decode('utf-8')
                return f"data:image/svg+xml;base64,{encoded}"

        except Exception as e:
            return None

    def get_file_display(self, file_path, tokens=None):
        """Returns HTML formatted string with icon and file info"""
        try:
            icon_data = self.get_file_icon(file_path)
            filename = Path(file_path).name

            if icon_data:
                display_text = (
                    f'<span class="file-item">'
                    f'<img src="{icon_data}" style="width:20px;height:20px;'
                    f'vertical-align:middle;margin-right:5px"/> {filename}'
                )
                if tokens is not None:
                    display_text += f' ({tokens:,} tokens)'
                display_text += '</span>'
                return display_text
            else:
                if tokens is not None:
                    return f"📄 {filename} ({tokens:,} tokens)"
                return f"📄 {filename}"

        except Exception as e:
            print(f"Error in get_file_display: {str(e)}")
            return Path(file_path).name
        

    def create_jira_story(self, story_content):
        import logging
        
        # Configure logging
        logging.basicConfig(level=logging.INFO)
        logger = logging.getLogger(__name__)
        
        # Add console handler if it doesn't exist
        if not logger.handlers:
            console_handler = logging.StreamHandler()
            console_handler.setLevel(logging.INFO)
            formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
            console_handler.setFormatter(formatter)
            logger.addHandler(console_handler)
        
        logger.info("\n=== Creating Jira Story ===")
        
        try:
            # Verify credentials
            email = os.getenv('CONFLUENCE_EMAIL')
            api_token = os.getenv('CONFLUENCE_API_KEY')
            
            if not email or not api_token:
                error_msg = "Missing Jira credentials"
                logger.error(error_msg)
                st.error(error_msg)
                return {"success": False, "error": error_msg}
            
            logger.info(f"Credentials check passed - Email: {email[:3]}...{email[-8:]}")
            
            # Validate story content
            logger.info("Validating story content...")
            required_fields = ['summary', 'description', 'priority', 'story_points']
            
            for field in required_fields:
                if field not in story_content:
                    error_msg = f"Missing required field: {field}"
                    logger.error(error_msg)
                    st.error(error_msg)
                    return {"success": False, "error": error_msg}
            
            logger.info("Story content validation passed")
            logger.info(f"Story content: {json.dumps(story_content, indent=2)}")
            
            # Prepare URL and headers
            url = "https://theseus-group.atlassian.net/rest/api/3/issue"
            logger.info(f"API URL: {url}")
            
            headers = {
                "Accept": "application/json",
                "Content-Type": "application/json"
            }
            
            # Prepare payload
            payload = {
                "fields": {
                    "project": {"key": "SCRUM"},
                    "summary": story_content["summary"],
                    "description": {
                        "type": "doc",
                        "version": 1,
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": story_content["description"]
                                    }
                                ]
                            }
                        ]
                    },
                    "issuetype": {"name": "Story"},
                    "priority": {"name": story_content["priority"]},
                    "customfield_10016": float(story_content["story_points"])
                }
            }
            
            # Add optional fields if they exist
            if "acceptance_criteria" in story_content:
                criteria_text = "\n\nAcceptance Criteria:\n" + "\n".join(f"- {c}" for c in story_content["acceptance_criteria"])
                payload["fields"]["description"]["content"].append({
                    "type": "paragraph",
                    "content": [{"type": "text", "text": criteria_text}]
                })
                
            if "technical_requirements" in story_content:
                reqs_text = "\n\nTechnical Requirements:\n" + "\n".join(f"- {r}" for r in story_content["technical_requirements"])
                payload["fields"]["description"]["content"].append({
                    "type": "paragraph",
                    "content": [{"type": "text", "text": reqs_text}]
                })
            
            logger.info("\nPayload to be sent:")
            logger.info(json.dumps(payload, indent=2))
            
            # Make API call
            logger.info("\nMaking API request...")
            try:
                response = requests.post(
                    url,
                    json=payload,
                    headers=headers,
                    auth=(email, api_token)
                )
                
                logger.info(f"Response Status: {response.status_code}")
                logger.info(f"Response Headers: {dict(response.headers)}")
                logger.info(f"Response Body: {response.text}")
                
                if response.status_code == 201:
                    issue_key = response.json()["key"]
                    success_msg = f"Story created successfully with key: {issue_key}"
                    logger.info(success_msg)
                    st.success(success_msg)
                    
                    # Add additional fields through updates if needed
                    if "estimation" in story_content:
                        logger.info(f"Adding estimation: {story_content['estimation']}")
                        # Code for updating estimation would go here
                    
                    if "risk_level" in story_content:
                        logger.info(f"Adding risk level: {story_content['risk_level']}")
                        # Code for updating risk level would go here
                    
                    return {"success": True, "issue_key": issue_key}
                else:
                    error_msg = f"Failed to create story. Status: {response.status_code}, Response: {response.text}"
                    logger.error(error_msg)
                    st.error(error_msg)
                    return {"success": False, "error": error_msg}
                    
            except requests.exceptions.RequestException as e:
                error_msg = f"API request failed: {str(e)}"
                logger.error(error_msg)
                st.error(error_msg)
                return {"success": False, "error": error_msg}
                
        except Exception as e:
            error_msg = f"Error creating story: {str(e)}"
            logger.error(error_msg)
            logger.error("Traceback:", exc_info=True)
            st.error(error_msg)
            return {"success": False, "error": error_msg}


    def generate_story_content(self, description, confluence_url=None):
        print("\n=== Starting Story Content Generation ===")
        print(f"Description: {description}")
        
        try:
            client = OpenAI()
            print("OpenAI client initialized")
            
            # First prompt for time, priority and risk analysis
            analysis_prompt = f"""
            Analyze the following user story description and determine:
            1. A realistic time estimation
            2. Priority level based on business impact and urgency
            3. Risk level based on technical complexity and dependencies

            Description:
            {description}

            IMPORTANT: Respond ONLY with a valid JSON object with this exact format:
            {{
                "estimation": "5d",
                "priority": "High/Medium/Low",
                "risk_level": "High/Medium/Low",
                "justification": {{
                    "time": "Time estimation explanation",
                    "priority": "Priority assignment explanation",
                    "risk": "Risk level explanation"
                }}
            }}
            """
            
            print("Making first API call for analysis...")
            analysis_response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert in user story analysis. Carefully evaluate each aspect and provide clear justifications."
                    },
                    {
                        "role": "user",
                        "content": analysis_prompt
                    }
                ],
                temperature=0.3
            )
            
            # Process analysis response
            analysis_text = analysis_response.choices[0].message.content.strip()
            print(f"Analysis response: {analysis_text}")
            
            try:
                analysis_data = json.loads(analysis_text)
                estimated_time = analysis_data.get("estimation", "5d")
                priority = analysis_data.get("priority", "Medium")
                risk_level = analysis_data.get("risk_level", "Medium")
                justifications = analysis_data.get("justification", {})
            except json.JSONDecodeError as e:
                print(f"Error parsing analysis response: {str(e)}")
                estimated_time = "5d"
                priority = "Medium"
                risk_level = "Medium"
                justifications = {}
            
            # Second prompt for complete story generation
            story_prompt = f"""
            Create a detailed Jira story based on this description and analysis:
            
            Description:
            {description}

            Analysis Results:
            - Estimation: {estimated_time}
            - Priority: {priority}
            - Risk Level: {risk_level}

            IMPORTANT: Respond ONLY with a JSON object that follows this exact structure:
            {{
                "summary": "Brief one-line summary",
                "description": "Detailed description including analyses justifications",
                "story_points": 8,
                "estimation": "{estimated_time}",
                "priority": "{priority}",
                "risk_level": "{risk_level}",
                "acceptance_criteria": [
                    "First acceptance criteria",
                    "Second acceptance criteria",
                    "Third acceptance criteria"
                ],
                "technical_requirements": [
                    "First technical requirement",
                    "Second technical requirement",
                    "Third technical requirement"
                ],
                "analysis_justification": {json.dumps(justifications)}
            }}
            """
            
            # Rest of story generation code...
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[{
                    "role": "system",
                    "content": "You are a Jira story generator. Always respond with valid JSON only."
                },
                {
                    "role": "user",
                    "content": story_prompt
                }],
                temperature=0.3
            )
            
            response_text = response.choices[0].message.content.strip()
            story_content = json.loads(response_text)
            
            # Ensure calculated values are maintained
            story_content['estimation'] = estimated_time
            story_content['priority'] = priority
            story_content['risk_level'] = risk_level
            
            # Validations
            if 'priority' not in story_content or story_content['priority'] not in ["High", "Medium", "Low"]:
                story_content['priority'] = "Medium"
                
            if 'risk_level' not in story_content or story_content['risk_level'] not in ["High", "Medium", "Low"]:
                story_content['risk_level'] = "Medium"
                
            return story_content
            
        except Exception as e:
            print(f"Error generating story content: {str(e)}")
            print("Traceback:", traceback.format_exc())
            return None